import os
import shutil
import sys
import subprocess

# 1. Install python-docx automatically if not present
try:
    import docx
except ImportError:
    print("Installing python-docx library for generating the Word report...")
    subprocess.run([sys.executable, "-m", "pip", "install", "python-docx"], check=True)
    import docx

from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

# 2. Define folders and file lists
src_dir = os.path.dirname(os.path.abspath(__file__))
sub_dir = os.path.join(src_dir, "submission")

print(f"Creating submission directory at: {sub_dir}")
os.makedirs(sub_dir, exist_ok=True)

files_to_copy = [
    # 01. CSV file
    "top300_1.csv",
    # 02. BED files
    "up_I1.bed", "up_I2.bed", "down_I1.bed", "down_I2.bed",
    # 03. FASTA files
    "up_I1.fa", "up_I2.fa", "down_I1.fa", "down_I2.fa",
    # 04. MEME results
    "up_I1_meme.txt", "up_I2_meme.txt", "down_I1_meme.txt", "down_I2_meme.txt",
    # Accompanying PDF reports for reference
    "I1_motif_comparison.png.pdf", "I2_motif_comparison.png.pdf"
]

# 3. Copy files to the submission folder
print("\nCopying files...")
for fname in files_to_copy:
    src_file = os.path.join(src_dir, fname)
    dest_file = os.path.join(sub_dir, fname)
    if os.path.exists(src_file):
        shutil.copy2(src_file, dest_file)
        print(f"Copied: {fname}")
    else:
        print(f"Warning: {fname} not found in workspace!")

# 4. Generate the Word Report (.docx) in the submission folder
print("\nGenerating motif_comparison_report.docx...")
doc = docx.Document()

# Adjust margins
sections = doc.sections
for section in sections:
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

# Add Document Title
title_p = doc.add_paragraph()
title_run = title_p.add_run("Motif Comparison Report")
title_run.font.size = Pt(24)
title_run.bold = True
title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Add Subtitle
sub_p = doc.add_paragraph()
sub_run = sub_p.add_run("Course: Introduction to Bioinformatics (MBIO0402), 2026\nProject: circRNA Flanking Intron Motif Discovery")
sub_run.font.size = Pt(11)
sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Spacer
doc.add_paragraph("")

# Section 1
doc.add_heading("1. Upstream Introns (I1) Comparison Summary (up_I1 vs down_I1)", level=1)
p1 = doc.add_paragraph()
p1.add_run("• Query Motif Set: ").bold = True
p1.add_run("up_I1_meme.txt (10 motifs)\n")
p1.add_run("• Target Motif Set: ").bold = True
p1.add_run("down_I1_meme.txt (10 motifs)\n")
p1.add_run("• Significance Threshold: ").bold = True
p1.add_run("E-value < 0.01\n")

doc.add_heading("Overlapping (Common) Motifs", level=2)
doc.add_paragraph(
    "Six out of the ten discovered motifs in up_I1 matched corresponding motifs in down_I1 below the significance threshold:"
)

# Table 1: Upstream Introns
t1 = doc.add_table(rows=1, cols=5)
t1.style = 'Table Grid'
hdr_cells = t1.rows[0].cells
hdr_cells[0].text = 'Upstream (UP) Motif ID'
hdr_cells[1].text = 'Downstream (DOWN) Motif ID'
hdr_cells[2].text = 'E-value'
hdr_cells[3].text = 'Best Alignment Consensus'
hdr_cells[4].text = 'Splicing Factor / Candidate'

# Header bolding and formatting
for cell in hdr_cells:
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.bold = True

rows_i1 = [
    ("MEME-3 (ARMCCTGKCT)", "MEME-6 (AGMCAGGGYT)", "9.51e-07", "ARMCCTGKCT vs AGMCAGGGYT", "GC-rich element (hnRNP / SR protein)"),
    ("MEME-1 (ACACACACAC)", "MEME-2 (TGKGTGTGTG)", "7.38e-05", "ACACACACAC vs TGKGTGTGTG", "CA/GT-repeats (CELF / MBNL family)"),
    ("MEME-2 (YTYTYTYTCY)", "MEME-5 (TCTYTCTYYC)", "2.16e-04", "YTYTYTYTCY vs TCTYTCTYYC", "Pyrimidine-rich tract (PTBP1 / hnRNP I)"),
    ("MEME-5 (TTTDTTTDTT)", "MEME-1 (TTTCTTTTTT)", "2.26e-04", "TTTDTTTDTT vs TTTCTTTTTT", "Poly-T tract (TDP-43 / hnRNP family)"),
    ("MEME-7 (CAGARGVCAG)", "MEME-7 (CTGRBCWCTS)", "3.82e-03", "CAGARGVCAG vs CTGRBCWCTS", "Purine-rich element (SRSF binding)"),
    ("MEME-10 (TGSYCTKCCW)", "MEME-9 (CTKCTCTTCY)", "6.13e-03", "TGSYCTKCCW vs CTKCTCTTCY", "Pyrimidine-rich motif (splicing regulator)")
]

for row in rows_i1:
    row_cells = t1.add_row().cells
    for i, item in enumerate(row):
        row_cells[i].text = item

doc.add_paragraph("")

doc.add_heading("Group-Specific Motifs", level=2)
doc.add_paragraph("Four motifs in up_I1 did not match any motif in down_I1 under the E-value threshold:")
p_spec1 = doc.add_paragraph()
p_spec1.add_run("UP-regulated specific motifs:\n").bold = True
p_spec1.add_run("• MEME-4 (YCCCAGCACY)\n• MEME-6 (TKGAACTCAS)\n• MEME-8 (CCAGRAGAGG)\n• MEME-9 (GCCACCATGT)")

doc.add_heading("Visual Alignment (I1)", level=2)
doc.add_paragraph("Refer to the uploaded alignment report: I1_motif_comparison.png.pdf")

doc.add_paragraph("")

# Section 2
doc.add_heading("2. Downstream Introns (I2) Comparison Summary (up_I2 vs down_I2)", level=1)
p2 = doc.add_paragraph()
p2.add_run("• Query Motif Set: ").bold = True
p2.add_run("up_I2_meme.txt (10 motifs)\n")
p2.add_run("• Target Motif Set: ").bold = True
p2.add_run("down_I2_meme.txt (10 motifs)\n")
p2.add_run("• Significance Threshold: ").bold = True
p2.add_run("E-value < 0.01\n")

doc.add_heading("Overlapping (Common) Motifs", level=2)
doc.add_paragraph(
    "Only two motifs (forming three matching pairs) in up_I2 matched corresponding motifs in down_I2 below the significance threshold:"
)

# Table 2: Downstream Introns
t2 = doc.add_table(rows=1, cols=5)
t2.style = 'Table Grid'
hdr_cells2 = t2.rows[0].cells
hdr_cells2[0].text = 'Upstream (UP) Motif ID'
hdr_cells2[1].text = 'Downstream (DOWN) Motif ID'
hdr_cells2[2].text = 'E-value'
hdr_cells2[3].text = 'Best Alignment Consensus'
hdr_cells2[4].text = 'Splicing Factor / Candidate'

# Header bolding and formatting
for cell in hdr_cells2:
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.bold = True

rows_i2 = [
    ("MEME-2 (CAMAMAMAMA)", "MEME-1 (TGTGTGTGTG)", "2.70e-04", "CAAACAAAAA vs TGTGTGTGTG", "CA/GT-repeats (CELF / MBNL family)"),
    ("MEME-2 (CAMAMAMAMA)", "MEME-4 (TTTTGTTTTT)", "1.34e-03", "CAAACAAAAA vs TTTTGTTTTT", "CA/T-rich elements (hnRNP family)"),
    ("MEME-5 (RAGTTCCAGG)", "MEME-9 (TBCTGGRAYT)", "1.39e-03", "GAGTTCCAGG vs TBCTGGRAYT", "Purine/Pyrimidine element (SRSF binding)")
]

for row in rows_i2:
    row_cells = t2.add_row().cells
    for i, item in enumerate(row):
        row_cells[i].text = item

doc.add_paragraph("")

doc.add_heading("Group-Specific Motifs", level=2)
doc.add_paragraph("Eight motifs in up_I2 did not match any motif in down_I2 under the E-value threshold:")
p_spec2 = doc.add_paragraph()
p_spec2.add_run("UP-regulated specific motifs:\n").bold = True
p_spec2.add_run("• MEME-1 (RGAGGMAGRG)\n• MEME-3 (AGCCWGGGCT)\n• MEME-4 (CCCAGCACHB)\n• MEME-6 (CAGCCTCCAG)\n• MEME-7 (SAGARGRMAG)\n• MEME-8 (GCABGCCTTT)\n• MEME-9 (GMAGGCAGAT)\n• MEME-10 (GCTTTGCCTT)")

doc.add_heading("Visual Alignment (I2)", level=2)
doc.add_paragraph("Refer to the uploaded alignment report: I2_motif_comparison.png.pdf")

doc.add_paragraph("")

# Section 3
doc.add_heading("3. Biological Conclusion", level=1)
p_conclusion1 = doc.add_paragraph()
r_c1 = p_conclusion1.add_run("1. Core Splicing Machinery Recognition: ")
r_c1.bold = True
p_conclusion1.add_run(
    "The overlapping motifs in both I1 and I2 regions are dominated by simple sequence repeats (SSRs), "
    "specifically CA-rich/GT-rich repeats (complementary to each other), pyrimidine-rich tracts (CT repeats), "
    "and poly-T tracts. These repeats serve as highly conserved binding platforms for core splicing regulators "
    "(such as CELF, PTBP1, and TDP-43 families). Their conservation across both UP and DOWN groups suggests they "
    "represent general structural elements required to bring flanking introns close together for back-splicing."
)

p_conclusion2 = doc.add_paragraph()
r_c2 = p_conclusion2.add_run("2. Dynamic Regulation via Group-Specific Motifs: ")
r_c2.bold = True
p_conclusion2.add_run(
    "The presence of numerous group-specific motifs (especially the 8 UP-regulated specific motifs in downstream "
    "introns) suggest that differential circRNA expression during cardiac pressure overload (TAC) is dynamically "
    "tuned by specific splicing enhancers or silencers (like SRSF proteins binding to purine-rich elements), "
    "rather than just the default back-splicing machinery."
)

report_path = os.path.join(sub_dir, "motif_comparison_report.docx")
doc.save(report_path)
print(f"Successfully generated Word report: {report_path}")
print("\nSubmission folder is ready!")
